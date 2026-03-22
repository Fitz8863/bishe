#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt
import copy

doc = Document('/tmp/paper_fixed.docx')

def find_para_index(doc, text):
    for i, para in enumerate(doc.paragraphs):
        if text in para.text:
            return i
    return -1

idx = find_para_index(doc, '3.2 YOLOv8 网络整体架构')
print(f"Found 3.2 section at paragraph index: {idx}")

if idx >= 0:
    reference = doc.paragraphs[idx]._p
    
    sections_content = [
        ('3.2.1 整体架构概述', 'heading3'),
        ('YOLOv8 是 Ultralytics 公司于 2023 年发布的最新一代 YOLO 系列目标检测算法，其网络架构延续了"骨干网络-颈部网络-检测头"（Backbone-Neck-Head）的经典设计范式，但在各模块的实现上进行了全面优化。', None),
        ('YOLOv8 的整体架构由三个主要部分组成：', None),
        ('Backbone（骨干网络）：采用 CSPDarknet，负责从输入图像中提取多尺度特征', 'bullet'),
        ('Neck（颈部网络）：采用 PANet 双向特征金字塔，实现多尺度特征融合', 'bullet'),
        ('Head（检测头）：采用无锚框解耦检测头，分别预测分类和边界框', 'bullet'),
        
        ('3.2.2 骨干网络（Backbone）', 'heading3'),
        ('YOLOv8 的骨干网络采用 CSPDarknet 作为特征提取器，由一系列卷积模块和 C2f 模块交替组成，负责从输入图像中逐层提取多尺度特征。', None),
        
        ('1. Focus 模块（空间下采样）', 'bold'),
        ('Focus 模块是 YOLOv8 对输入图像进行空间降采样的关键组件，其核心思想是通过像素切片操作将特征图的空间维度减半，同时通道数扩展至原来的 4 倍。', None),
        ('数学定义：', 'bold'),
        ('对于输入特征图 X ∈ R^(H×W×C)，Focus 模块将其空间下采样 2×，输出 Y ∈ R^((H/2)×(W/2)×4C)：', None),
        ('Y(i,j,k) = X(row(k), col(k))', None),
        ('其中：row(k) = ⌊k/4⌋ mod H，col(k) = k mod W', None),
        ('Focus 模块不包含可学习参数，仅进行通道重组操作。', None),
        
        ('2. C2f 模块（CSP Bottleneck with 2 Convolutions）', 'bold'),
        ('C2f 模块是 YOLOv8 的核心构建块，源自 YOLOv5 的 C3 模块并进行了改进。其设计融合了跨阶段局部网络（CSPNet）的思想，通过特征分离和密集连接增强梯度流动。', None),
        ('数学定义：', 'bold'),
        ('给定输入特征 x ∈ R^(H×W×C_in)，C2f 模块的计算过程如下：', None),
        ('x₁ = Conv₁(x)    # 1×1 卷积，通道分离', None),
        ('x₂ = Split(x₁)   # 通道分割', None),
        ('x₃ = Bottleneck₁(x₂)  # 瓶颈层', None),
        ('x₄ = Bottleneck₂(x₃)  # 瓶颈层', None),
        ('...', None),
        ('xₙ₊₂ = Bottleneckₙ(xₙ₊₁)', None),
        ('y = Conv₂(Concat(x₂, x₃, ..., xₙ₊₂))', None),
        ('其中：', 'bold'),
        ('• Conv₁：1×1 卷积，用于通道分离', None),
        ('• Split：通道分割操作，将特征分为两部分', None),
        ('• Bottleneck：瓶颈层，由 1×1 卷积 + 3×3 卷积 + 残差连接组成', None),
        ('• Concat：通道维度拼接', None),
        ('• Conv₂：1×1 卷积，用于通道融合', None),
        
        ('3. SPPF 模块（Spatial Pyramid Pooling - Fast）', 'bold'),
        ('SPPF 模块通过多尺度最大池化操作实现空间金字塔池化，有效融合不同感受野的特征，增强模型对目标尺度变化的鲁棒性。', None),
        ('数学定义：', 'bold'),
        ('对于输入特征 x，SPPF 模块的计算过程为：', None),
        ('y = Concat(MaxPool(x, k=5), MaxPool(x₁, k=5), MaxPool(x₂, k=5), x)', None),
        ('其中 x₁ = MaxPool(x, k=5)，x₂ = MaxPool(x₁, k=5)。', None),
        
        ('3.2.3 颈部网络（Neck）', 'heading3'),
        ('YOLOv8 的颈部网络采用路径聚合网络（PANet）结合特征金字塔网络（FPN）的双向特征融合架构，实现多尺度特征的充分交互。', None),
        ('PANet 在 FPN 自顶向下融合的基础上，额外增加了自底向上的特征增强路径，形成双向特征流动。', None),
        ('数学定义：', 'bold'),
        ('设 Pᵢ 表示 FPN 第 i 层的特征，Nᵢ 表示 PAN 第 i 层的增强特征：', None),
        ('自顶向下融合（FPN）：Pᵢ = Conv(Upsample(P_(i-1))) ⊕ Conv(Cᵢ)', None),
        ('自底向上增强（PAN）：Nᵢ = Conv(Concat(Pᵢ, Conv(Downsample(N_(i-1)))))', None),
        
        ('3.2.4 检测头（Head）', 'heading3'),
        
        ('1. 解耦头设计', 'bold'),
        ('YOLOv8 将分类和回归任务完全解耦，使用独立的分支进行预测。', None),
        ('• 分类分支：f_cls = Sigmoid(Conv₁ₓ₁(Conv₃ₓ₃(F)))', None),
        ('• 回归分支：b_box = Conv₁ₓ₁(Conv₃ₓ₃(F))', None),
        
        ('2. 多尺度检测', 'bold'),
        ('YOLOv8 在三个不同尺度的特征图上进行检测：', None),
        ('• P3 (80×80)：检测大目标，如香烟、打火机', None),
        ('• P4 (40×40)：检测中目标，如未穿防护服人员', None),
        ('• P5 (20×20)：检测小目标，如安全帽、口罩', None),
        
        ('3. 无锚框机制', 'bold'),
        ('YOLOv8 舍弃了基于锚框（Anchor-based）的检测范式，采用无锚框（Anchor-free）策略。', None),
        ('预测输出：o(i,j) = [d_x, d_y, d_w, d_h, p_obj, p₁, p₂, ..., p_C]', None),
        ('其中：d_x, d_y 为边界框中心偏移量；d_w, d_h 为宽高的自然对数；p_obj 为目标置信度；p₁~p_C 为类别概率。', None),
        
        ('3.2.5 损失函数', 'heading3'),
        ('YOLOv8 的总损失由三部分组成：分类损失、回归损失和置信度损失。', None),
        ('1. 分类损失（BCE Loss）：', 'bold'),
        ('L_cls = -1/N Σ [y_i · log(ŷ_i) + (1-y_i) · log(1-ŷ_i)]', None),
        ('2. 回归损失（CIoU + DFL Loss）：', 'bold'),
        ('L_CIoU = 1 - IoU + ρ²(b, b^gt)/c² + αv', None),
        ('其中 IoU 为预测框与真实框的交并比；ρ²(b, b^gt) 为中心点欧氏距离；c 为最小闭包对角线长度；αv 为长宽比因子。', None),
        ('3. 总损失：', 'bold'),
        ('L_total = L_cls + λ₁·L_CIoU + λ₂·L_DFL + λ₃·L_obj', None),
        ('其中 λ₁=7.5, λ₂=1.5, λ₃=1.0。', None),
        
        ('3.2.6 本系统的 YOLOv8 适配', 'heading3'),
        ('针对化工厂安全行为监控的实际需求，本系统对标准 YOLOv8 进行了针对性适配：', None),
        ('• 输入尺寸：640×640（平衡精度与推理速度）', None),
        ('• 类别数：5（吸烟、动火、无防护装备1/2/3）', None),
        ('• 骨干网络：YOLOv8s（在 Jetson Orin Nano 上的性能平衡）', None),
        ('• 检测尺度：80×80, 40×40, 20×20（多尺度覆盖不同大小目标）', None),
    ]
    
    for text, style in sections_content:
        new_para = copy.deepcopy(reference)
        ref_idx = list(reference.getparent()).index(reference)
        reference.addnext(new_para)
        reference = new_para
        
        para = doc.paragraphs[ref_idx + 1]
        para.clear()
        run = para.add_run(text)
        
        if style == 'heading3':
            run.bold = True
            run.font.size = Pt(14)
        elif style == 'bold':
            run.bold = True
        elif style == 'bullet':
            para.style = doc.styles['List Bullet']

    output_path = '/home/hwj/jetson/bishe/论文.docx'
    doc.save(output_path)
    print(f'Done: {output_path}')
else:
    print('Section not found!')
