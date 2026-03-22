#!/usr/bin/env python3
"""Generate 3.2 section content as docx"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

styles = doc.styles

para = doc.add_paragraph()
run = para.add_run('3.2 YOLOv8 网络整体架构')
run.bold = True
run.font.size = Pt(14)

doc.add_heading('3.2.1 整体架构概述', level=3)

doc.add_paragraph(
    'YOLOv8 是 Ultralytics 公司于 2023 年发布的最新一代 YOLO 系列目标检测算法，其网络架构延续了"骨干网络-颈部网络-检测头"（Backbone-Neck-Head）的经典设计范式，但在各模块的实现上进行了全面优化。'
)

p = doc.add_paragraph()
p.add_run('图 3-2 YOLOv8 网络整体架构图').italic = True

doc.add_paragraph(
    'YOLOv8 的整体架构由三个主要部分组成：'
)
doc.add_paragraph('Backbone（骨干网络）：采用 CSPDarknet，负责从输入图像中提取多尺度特征', style='List Bullet')
doc.add_paragraph('Neck（颈部网络）：采用 PANet 双向特征金字塔，实现多尺度特征融合', style='List Bullet')
doc.add_paragraph('Head（检测头）：采用无锚框解耦检测头，分别预测分类和边界框', style='List Bullet')

doc.add_heading('3.2.2 骨干网络（Backbone）', level=3)

doc.add_paragraph(
    'YOLOv8 的骨干网络采用 CSPDarknet 作为特征提取器，由一系列卷积模块和 C2f 模块交替组成，负责从输入图像中逐层提取多尺度特征。'
)

h = doc.add_heading('1. Focus 模块（空间下采样）', level=4)
p = doc.add_paragraph()
p.add_run('Focus 模块是 YOLOv8 对输入图像进行空间降采样的关键组件，其核心思想是通过像素切片（Pixel Space Sampling）操作将特征图的空间维度减半，同时通道数扩展至原来的 4 倍。').bold = False

p = doc.add_paragraph()
p.add_run('数学定义：').bold = True

doc.add_paragraph(
    '对于输入特征图 X ∈ R^(H×W×C)，Focus 模块将其空间下采样 2×，输出 Y ∈ R^((H/2)×(W/2)×4C)：\n\n'
    'Y(i,j,k) = X(row(k), col(k))\n\n'
    '其中：\n'
    '• row(k) = ⌊k/4⌋ mod H\n'
    '• col(k) = k mod W'
)

h = doc.add_heading('2. C2f 模块（CSP Bottleneck with 2 Convolutions）', level=4)
doc.add_paragraph(
    'C2f 模块是 YOLOv8 的核心构建块，源自 YOLOv5 的 C3 模块并进行了改进。其设计融合了跨阶段局部网络（CSPNet）的思想，通过特征分离和密集连接增强梯度流动。'
)

p = doc.add_paragraph()
p.add_run('数学定义：').bold = True

doc.add_paragraph(
    '给定输入特征 x ∈ R^(H×W×C_in)，C2f 模块的计算过程如下：\n\n'
    'x₁ = Conv₁(x)                    # 1×1 卷积，通道分离\n'
    'x₂ = Split(x₁)                   # 通道分割\n'
    'x₃ = Bottleneck₁(x₂)             # 瓶颈层\n'
    'x₄ = Bottleneck₂(x₃)             # 瓶颈层\n'
    'x₅ = Bottleneck₃(x₄)             # 瓶颈层\n'
    '...\n'
    'xₙ₊₂ = Bottleneckₙ(xₙ₊₁)\n'
    'y = Conv₂(Concat(x₂, x₃, x₄, ..., xₙ₊₂))  # 通道拼接后 1×1 卷积'
)

p = doc.add_paragraph()
p.add_run('其中：').bold = True
doc.add_paragraph('• Conv₁：1×1 卷积，用于通道分离', style='List Bullet')
doc.add_paragraph('• Split：通道分割操作，将特征分为两部分', style='List Bullet')
doc.add_paragraph('• Bottleneck：瓶颈层，由 1×1 卷积 + 3×3 卷积 + 残差连接组成', style='List Bullet')
doc.add_paragraph('• Concat：通道维度拼接', style='List Bullet')
doc.add_paragraph('• Conv₂：1×1 卷积，用于通道融合', style='List Bullet')

h = doc.add_heading('3. SPPF 模块（Spatial Pyramid Pooling - Fast）', level=4)
doc.add_paragraph(
    'SPPF 模块通过多尺度最大池化操作实现空间金字塔池化，有效融合不同感受野的特征，增强模型对目标尺度变化的鲁棒性。'
)

p = doc.add_paragraph()
p.add_run('数学定义：').bold = True

doc.add_paragraph(
    '对于输入特征 x，SPPF 模块的计算过程为：\n\n'
    'y = Concat(MaxPool(x, k=5), MaxPool(x₁, k=5), MaxPool(x₂, k=5), x)\n\n'
    '其中 x₁ = MaxPool(x, k=5)，x₂ = MaxPool(x₁, k=5)。'
)

doc.add_heading('3.2.3 颈部网络（Neck）', level=3)

doc.add_paragraph(
    'YOLOv8 的颈部网络采用路径聚合网络（PANet）结合特征金字塔网络（FPN）的双向特征融合架构，实现多尺度特征的充分交互。'
)

p = doc.add_paragraph()
p.add_run('PANet 双向特征融合').bold = True

doc.add_paragraph(
    'PANet 在 FPN 自顶向下融合的基础上，额外增加了自底向上的特征增强路径，形成双向特征流动。'
)

p = doc.add_paragraph()
p.add_run('数学定义：').bold = True

doc.add_paragraph(
    '设 P_i 表示 FPN 第 i 层的特征，N_i 表示 PAN 第 i 层的增强特征：\n\n'
    '自顶向下融合（FPN）：\n'
    'P_i = Conv(Upsample(P_(i-1))) ⊕ Conv(C_i)\n\n'
    '自底向上增强（PAN）：\n'
    'N_i = Conv(Concat(P_i, Conv(Downsample(N_(i-1)))))'
)

doc.add_heading('3.2.4 检测头（Head）', level=3)

p = doc.add_paragraph()
p.add_run('1. 解耦头设计').bold = True

doc.add_paragraph(
    'YOLOv8 将分类和回归任务完全解耦，使用独立的分支进行预测。'
)

p = doc.add_paragraph()
p.add_run('分类分支：').bold = True
doc.add_paragraph('f_cls = Conv(3×3)(F)\nf_cls = Conv(1×1)(f_cls)\np_cls = Sigmoid(f_cls)')

p = doc.add_paragraph()
p.add_run('回归分支：').bold = True
doc.add_paragraph('f_reg = Conv(3×3)(F)\nf_reg = Conv(1×1)(f_reg)\nb_box = f_reg')

p = doc.add_paragraph()
p.add_run('2. 多尺度检测').bold = True

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '检测层'
hdr[1].text = '特征图尺寸'
hdr[2].text = '感受野'
hdr[3].text = '适用目标'
row1 = table.rows[1].cells
row1[0].text = 'P3 (大目标)'
row1[1].text = '80×80'
row1[2].text = '小'
row1[3].text = '香烟、打火机'
row2 = table.rows[2].cells
row2[0].text = 'P4 (中目标)'
row2[1].text = '40×40'
row2[2].text = '中'
row2[3].text = '未穿防护服人员'
row3 = table.rows[3].cells
row3[0].text = 'P5 (小目标)'
row3[1].text = '20×20'
row3[2].text = '大'
row3[3].text = '安全帽、口罩'

p = doc.add_paragraph()
p.add_run('3. 无锚框机制').bold = True

doc.add_paragraph(
    'YOLOv8 舍弃了基于锚框（Anchor-based）的检测范式，采用无锚框（Anchor-free）策略，直接预测目标边界框的四个边界相对于特征图网格点的偏移量。'
)

doc.add_paragraph(
    '对于特征图上每个位置 (i, j)，预测输出为：\n\n'
    'o(i,j) = [d_x, d_y, d_w, d_h, p_obj, p₁, p₂, ..., p_C]\n\n'
    '其中：\n'
    '• d_x, d_y：边界框中心相对于网格左上角的偏移量\n'
    '• d_w, d_h：边界框宽高的自然对数\n'
    '• p_obj：目标置信度\n'
    '• p₁, p₂, ..., p_C：各类别的分类概率'
)

doc.add_heading('3.2.5 损失函数', level=3)

p = doc.add_paragraph()
p.add_run('1. 分类损失（BCE Loss）').bold = True

doc.add_paragraph(
    'L_cls = -1/N Σ [y_i · log(ŷ_i) + (1-y_i) · log(1-ŷ_i)]'
)

p = doc.add_paragraph()
p.add_run('2. 回归损失（CIoU + DFL Loss）').bold = True

doc.add_paragraph(
    'L_CIoU = 1 - IoU + ρ²(b, b^gt)/c² + αv\n\n'
    '其中：\n'
    '• IoU：预测框与真实框的交并比\n'
    '• ρ²(b, b^gt)：两个框中心点的欧氏距离\n'
    '• c：覆盖两个框的最小闭包区域的对角线长度\n'
    '• α：权重因子\n'
    '• v：长宽比一致性度量'
)

p = doc.add_paragraph()
p.add_run('3. 总损失').bold = True

doc.add_paragraph(
    'L_total = L_cls + λ₁·L_CIoU + λ₂·L_DFL + λ₃·L_obj\n\n'
    '其中 λ₁, λ₂, λ₃ 为各损失项的权重系数（默认为 λ₁=7.5, λ₂=1.5, λ₃=1.0）。'
)

doc.add_heading('3.2.6 本系统的 YOLOv8 适配', level=3)

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '参数'
hdr[1].text = '配置值'
hdr[2].text = '说明'
row1 = table.rows[1].cells
row1[0].text = '输入尺寸'
row1[1].text = '640×640'
row1[2].text = '平衡精度与推理速度'
row2 = table.rows[2].cells
row2[0].text = '类别数'
row2[1].text = '5'
row2[2].text = '吸烟、动火、无防护装备1/2/3'
row3 = table.rows[3].cells
row3[0].text = '骨干网络'
row3[1].text = 'YOLOv8s'
row3[2].text = 'Jetson Orin Nano 性能平衡'
row4 = table.rows[4].cells
row4[0].text = '检测尺度'
row4[1].text = '80×80, 40×40, 20×20'
row4[2].text = '多尺度覆盖'

doc.save('/home/hwj/jetson/bishe/section_3_2_content.docx')
print('Done: section_3_2_content.docx')
